import openpyxl as excel
import docx,random,os,shutil
import english_sort
import copy

create_file = 1             #生成するファイル数
random_list = []            #出題する問題を格納
section_tuple = (1,3,4,5,6) #問題の章番号
question_num  = (6,6,6,6,6) #各章からの出題数

##変数の定義
def define():
    global doc,ans_doc,wb,sheet,question
    english_sort.sort()
    doc = docx.Document('exampaper_template_englishb.docx')             #出力用word文書テンプレート
    ans_doc = docx.Document('exampaper_answer_template_englishb.docx')  #解答出力用word文書テンプレート
    wb = excel.load_workbook('databank_englishB-copy.xlsx')             #入力データ
    sheet = wb['シート1']
    question = []

##問題をリストに代入
def set_question(sheet_local):
    for s in section_tuple:
        section_list = []
        for i in sheet_local.iter_rows():
            if i[0].value == s:
                section_list.append([i[1].value,i[2].value,i[3].value,i[0].value,i[4].value])
        question.append(section_list)
    return None

##出題する問題を設定
def choice_question():
    random_list = []
    for set_question in range(len(section_tuple)):
        random_list += random.sample(question[set_question],question_num[set_question])
    return random_list

##word文書に問題用紙を出力
def write_question(random_list):
    for para in doc.paragraphs:
        if 'japanese' in para.text:
            para.text = para.text.replace('japanese',str(random_list[0][0]))
        if 'english' in para.text:
            para.text = para.text.replace('english',str(random_list[0][2]))
            del random_list[0]
    return None

##word文書に解答用紙を出力
def write_answer(answer_list):
    current = len(random_list)
    japanese = len(random_list)
    keep_list = []
    keep_list.append(copy.deepcopy(random_list))
    for para in ans_doc.paragraphs:
        try:
            if 'japanese' in para.text:
                japanese -= 1
                para.text = para.text.replace('japanese',str(random_list[0][0]))
            if 'english' in para.text:
                para.text = para.text.replace('english',str(random_list[0][2]))
            if 'answer' in para.text:
                current -= 1
                if current != japanese:
                    print(keep_list[0][30-len(random_list)][4])
                para.text = para.text.replace('answer',str(random_list[0][1]))
                del random_list[0]
        except IndexError:
            pass
    return None

##word文書の書き出し
def output(name,file_num):
    if not os.path.isdir('exampaper'):
        os.makedirs('exampaper')
        print("make directory 'exampaper' completed")
    while(1):
        file_name = 'exampaper_englishb'+str(file_num)+name+'.docx'
        if not os.path.isfile('exampaper/'+file_name):
            break;
        file_num += 1
    if name == '':
        doc.save(file_name)
    else :
        ans_doc.save(file_name)
    shutil.move(file_name,'exampaper')
    print("save '" + file_name + "' completed")
    return file_num

##処理部分
file_num = 1
for i in range(create_file):
    define()
    set_question(sheet)
    random_list = choice_question()
    random_list = random_list
    write_question(random_list)
    write_answer(random_list)
    file_num = output('',file_num)
    output('_answer',file_num)
print('all process completed')